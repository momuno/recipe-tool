"""
Processes any file conversions, whether for references files or generated files.
"""

import os
import pypandoc
from docx import Document


def markdown_to_docx(markdown_content: str, output_path: str) -> str:
    """Convert markdown content to docx file and return the output path."""
    try:
        pypandoc.convert_text(
            markdown_content,
            "docx",
            format="md",
            outputfile=output_path,
            extra_args=["--extract-media=."],  # Extract images to current directory
        )
        return output_path
    except Exception as e:
        raise Exception(f"Failed to convert markdown to docx: {str(e)}")


def check_docx_protected(docx_path: str) -> tuple[bool, str]:
    """Check if a docx file is protected/encrypted without fully extracting text.
    Returns (is_protected, error_message)
    """
    try:
        from docx import Document

        filename = os.path.basename(docx_path)
        # Try to open the document
        doc = Document(docx_path)
        # Try to access at least one paragraph to ensure it's readable
        _ = len(doc.paragraphs)
        return False, ""
    except Exception as e:
        error_msg = str(e).lower()
        filename = os.path.basename(docx_path)
        # Check for common security/encryption error messages
        if any(term in error_msg for term in ["package not found"]):
            return True, (f"Document '{filename}' appears to be protected or encrypted and cannot be processed.")
        else:
            # Some other error, but not protection-related
            return False, f"Document '{filename}' may have issues: {str(e)}"


def docx_to_text(docx_path: str) -> str:
    """Extract text content from a docx file."""
    try:
        doc = Document(docx_path)
        paragraphs = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                paragraphs.append(paragraph.text.strip())

        return "\n\n".join(paragraphs)
    except Exception as e:
        error_msg = str(e).lower()
        filename = os.path.basename(docx_path)
        # Check for common security/encryption error messages
        if any(term in error_msg for term in ["package not found"]):
            raise Exception(f"Document '{filename}' may be protected or encrypted and cannot be processed.")
        else:
            raise Exception(f"Failed to extract text from '{filename}': {str(e)}")
